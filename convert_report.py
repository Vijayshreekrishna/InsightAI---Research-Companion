
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_formatting(paragraph, text):
    # Handle bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic *text* inside non-bold parts
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*') and len(sub) > 2:
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub)

def create_table(doc, lines):
    # Parse header
    header_line = lines[0]
    headers = [h.strip() for h in header_line.split('|') if h.strip()]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h.replace('**', '') # Remove bold from header if present, usually auto-bolded by style

    # Parse rows, skipping separator line if present (contains ---)
    for line in lines[1:]:
        if '---' in line:
            continue
        cols = [c.strip() for c in line.split('|')][1:-1] # split and remove first/last empty due to | border
        if not cols: continue
        
        row_cells = table.add_row().cells
        for i, val in enumerate(cols):
            if i < len(row_cells):
                row_cells[i].text = val.replace('**', '') # simplistic content

def convert_md_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_buffer = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Handle Table
        if line.startswith('|'):
            table_buffer.append(line)
            in_table = True
            continue
        else:
            if in_table:
                create_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if not line:
            continue

        # Page Break
        if '<div style="page-break-after: always;"></div>' in line:
            doc.add_page_break()
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            apply_formatting(p, line[2:])
        
        # Images
        elif line.startswith('!['):
            # Just put text placeholder
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt, url = match.groups()
                run = p.add_run(f"[IMAGE: {alt}]")
                run.italic = True
                run.font.color.rgb = RGBColor(255, 0, 0) # Red placeholder
            else:
                p.add_run(line)

        # Normal Text
        else:
            p = doc.add_paragraph()
            if line.startswith('<br>'):
                continue # Skip standalone br
            apply_formatting(p, line)

    if in_table:
        create_table(doc, table_buffer)

    doc.save(docx_file)
    print(f"Successfully created {docx_file}")

if __name__ == "__main__":
    md_path = r"c:\Users\vijay\.gemini\antigravity\brain\ad21e52b-11aa-4d38-8eea-2e0de81edc20\Insight_AI_Phase1_Report.md"
    docx_path = r"c:\Users\vijay\Downloads\vsk\gdg\Insight_AI_Phase1_Report.docx"
    convert_md_to_docx(md_path, docx_path)
