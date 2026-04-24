# utils/paper_utils.py
import os
import json
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

HISTORY_FILE = "paper_history.json"

class PDF(FPDF):
    def header(self):
        pass
    def footer(self):
        self.set_y(-15)
        self.set_font("helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

def get_template_config(template_name):
    """Return styling configuration for different academic templates."""
    configs = {
        "IEEE": {"font_family": "Times", "title_size": 24, "section_size": 11, "body_size": 10, "columns": 2, "line_spacing": 1.0},
        "ACM": {"font_family": "Helvetica", "title_size": 18, "section_size": 12, "body_size": 9, "columns": 1, "line_spacing": 1.15},
        "Springer": {"font_family": "Times", "title_size": 16, "section_size": 11, "body_size": 10, "columns": 1, "line_spacing": 1.2},
        "Generic University": {"font_family": "Arial", "title_size": 16, "section_size": 14, "body_size": 12, "columns": 1, "line_spacing": 1.5}
    }
    return configs.get(template_name, configs["Generic University"])

def clean_text_for_pdf(text: str) -> str:
    if not text: return ""
    replacements = {"√": "sqrt", "→": "->", "←": "<-", "≥": ">=", "≤": "<=", "≈": "~", "•": "*", "…": "...", "—": "-", "–": "-", "°": " deg", "±": "+/-"}
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text.encode('latin-1', 'replace').decode('latin-1')

def export_as_pdf(paper_data, template_name="Generic University", font_style="Serif", line_spacing=1.15):
    config = get_template_config(template_name)
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=False)
    pdf.add_page()
    
    font_map = {"Times New Roman": "Times", "Arial": "Helvetica", "Calibri": "Helvetica", "Georgia": "Times", "Cambria": "Times", "Monospace": "Courier"}
    pdf_font = font_map.get(font_style, "Times")
    is_multi_col = template_name in ["IEEE", "ACM"]
    lh = 6 * line_spacing
    margin, col_gap = 15, 10
    col_width = (pdf.w - (2 * margin) - col_gap) / 2
    y_start_top = 25

    # 1. Title
    pdf.set_font(pdf_font, "B", config["title_size"])
    pdf.multi_cell(0, 10, clean_text_for_pdf(paper_data.get("title", "Research Paper")), align="C")
    pdf.ln(8)
    
    # 2. Abstract
    pdf.set_font(pdf_font, "B", config["body_size"])
    pdf.write(lh, "ABSTRACT - ")
    pdf.set_font(pdf_font, "I", config["body_size"])
    pdf.multi_cell(0, lh, clean_text_for_pdf(str(paper_data.get("abstract", ""))), align="J")
    pdf.ln(3)
    
    # 3. Keywords
    pdf.set_font(pdf_font, "B", config["body_size"])
    pdf.write(lh, "KEYWORDS: ")
    pdf.set_font(pdf_font, "", config["body_size"])
    pdf.multi_cell(0, lh, clean_text_for_pdf(str(paper_data.get("keywords", ""))), align="L")
    pdf.ln(12)

    y_after_header = pdf.get_y()
    sections = [("Introduction", "introduction"), ("Literature Review", "lit_review"), ("Methodology", "methodology"), ("Results & Discussion", "results"), ("Conclusion", "conclusion"), ("Future Scope", "future_scope"), ("References", "references")]

    current_col = 0
    def switch_col_or_page():
        nonlocal current_col
        if current_col == 0:
            current_col = 1
            target_y = y_after_header if pdf.page == 1 else y_start_top
            pdf.set_xy(margin + col_width + col_gap, target_y)
        else:
            pdf.add_page()
            current_col = 0
            pdf.set_xy(margin, y_start_top)

    if is_multi_col:
        for label, key in sections:
            content = paper_data.get(key, "")
            if not content: continue
            if pdf.get_y() > 260: switch_col_or_page()
            pdf.set_font(pdf_font, "B", config["section_size"])
            pdf.set_x(margin if current_col == 0 else margin + col_width + col_gap)
            pdf.cell(col_width, 10, label.upper(), ln=True, border="B")
            pdf.ln(4)
            pdf.set_font(pdf_font, "", config["body_size"])
            paras = str(content).split("\n\n")
            for p in paras:
                if not p.strip(): continue
                lines_list = pdf.multi_cell(col_width, lh, clean_text_for_pdf(p.strip()), align="J", split_only=True)
                for i, line in enumerate(lines_list):
                    if pdf.get_y() > 275: switch_col_or_page()
                    pdf.set_x(margin if current_col == 0 else margin + col_width + col_gap)
                    # FIX: Use multi_cell for single line to allow justification
                    pdf.multi_cell(col_width, lh, line, align="J" if i < len(lines_list) - 1 else "L")
                pdf.ln(5)
    else:
        pdf.set_auto_page_break(auto=True, margin=20)
        for label, key in sections:
            content = paper_data.get(key, "")
            if not content: continue
            pdf.set_font(pdf_font, "B", config["section_size"])
            pdf.set_fill_color(245, 245, 245)
            pdf.cell(0, 10, "  " + label.upper(), ln=True, fill=True)
            pdf.ln(4)
            pdf.set_font(pdf_font, "", config["body_size"])
            pdf.multi_cell(0, lh + 1, clean_text_for_pdf(str(content)), align="J")
            pdf.ln(8)
    return bytes(pdf.output())

def export_as_docx(paper_data, template_name="Generic University", font_style="Serif", line_spacing=1.15):
    config = get_template_config(template_name)
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.75)
    font_map = {"Times New Roman": "Times New Roman", "Arial": "Arial", "Calibri": "Calibri", "Georgia": "Georgia", "Cambria": "Cambria", "Monospace": "Courier New"}
    word_font = font_map.get(font_style, "Times New Roman")
    is_multi_col = template_name in ["IEEE", "ACM"]

    title = doc.add_heading(paper_data.get("title", "Research Paper"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name, run.font.size, run.font.bold, run.font.color.rgb = word_font, Pt(config["title_size"]), True, RGBColor(0, 0, 0)

    abs_para = doc.add_paragraph()
    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = abs_para.add_run("Abstract— ")
    run_label.bold = True
    run_label.font.name = word_font
    abs_text = abs_para.add_run(str(paper_data.get("abstract", "")))
    abs_text.font.name, abs_text.italic = word_font, True

    if is_multi_col:
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        sectPr = new_section._sectPr
        cols = sectPr.xpath("./w:cols")[0]
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "400")
    
    sections = [("Keywords", "keywords"), ("Introduction", "introduction"), ("Literature Review", "lit_review"), ("Methodology", "methodology"), ("Results & Discussion", "results"), ("Conclusion", "conclusion"), ("Future Scope", "future_scope"), ("References", "references")]
    for label, key in sections:
        content = paper_data.get(key, "")
        if not content: continue
        h = doc.add_heading(label.upper(), level=1)
        for run in h.runs:
            run.font.name, run.font.size, run.font.bold, run.font.color.rgb = word_font, Pt(config["section_size"]), True, RGBColor(0, 0, 0)
        paras = str(content).split("\n\n")
        for p_text in paras:
            if not p_text.strip(): continue
            p = doc.add_paragraph(p_text.strip())
            p.alignment, p.paragraph_format.line_spacing, p.paragraph_format.first_line_indent = WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing, Inches(0.2)
            for run in p.runs:
                run.font.name, run.font.size = word_font, Pt(config["body_size"])
    import io
    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

def save_to_history(paper_data):
    history = load_history()
    import datetime
    if "timestamp" not in paper_data:
        paper_data["timestamp"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    history.insert(0, paper_data)
    history = history[:20]
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=4)

def load_history():
    if not os.path.exists(HISTORY_FILE): return []
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f: return json.load(f)
    except: return []

def delete_from_history(index):
    history = load_history()
    if 0 <= index < len(history):
        history.pop(index)
        with open(HISTORY_FILE, "w", encoding="utf-8") as f: json.dump(history, f, indent=4)
        return True
    return False
